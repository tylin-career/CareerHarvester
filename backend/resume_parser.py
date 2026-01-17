"""
Resume Parser Module
--------------------
Provides robust text extraction from various resume file formats (PDF, DOCX, TXT).
Uses PyPDF2 for PDFs and python-docx for Word documents.
"""

import io
import logging
from typing import Optional

# Configure logging
logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_content: Raw bytes of the PDF file
        
    Returns:
        Extracted text content as a string
    """
    try:
        from PyPDF2 import PdfReader
        
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        full_text = "\n\n".join(text_parts)
        
        if not full_text.strip():
            logger.warning("PDF appears to be empty or contains only images")
            return "[PDF contains no extractable text - may be image-based]"
            
        return full_text.strip()
        
    except ImportError:
        logger.error("PyPDF2 not installed. Run: pip install PyPDF2")
        raise RuntimeError("PDF parsing library not available")
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        raise RuntimeError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_content: Raw bytes of the DOCX file
        
    Returns:
        Extracted text content as a string
    """
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)
        
        if not full_text.strip():
            logger.warning("DOCX appears to be empty")
            return "[Document contains no extractable text]"
            
        return full_text.strip()
        
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise RuntimeError("DOCX parsing library not available")
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise RuntimeError(f"Failed to parse DOCX: {str(e)}")


def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text from plain text file with encoding detection.
    
    Args:
        file_content: Raw bytes of the text file
        
    Returns:
        Decoded text content
    """
    # Try common encodings
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            return file_content.decode(encoding).strip()
        except (UnicodeDecodeError, LookupError):
            continue
    
    # Last resort: decode with replacement
    return file_content.decode('utf-8', errors='replace').strip()


def parse_resume(file_content: bytes, filename: str, mime_type: Optional[str] = None) -> str:
    """
    Parse resume content from various file formats.
    
    Args:
        file_content: Raw bytes of the uploaded file
        filename: Original filename (used for extension detection)
        mime_type: Optional MIME type for format detection
        
    Returns:
        Extracted text content from the resume
        
    Raises:
        ValueError: If file format is not supported
        RuntimeError: If parsing fails
    """
    filename_lower = filename.lower()
    
    # Determine file type
    if filename_lower.endswith('.pdf') or mime_type == 'application/pdf':
        logger.info(f"Parsing PDF resume: {filename}")
        return extract_text_from_pdf(file_content)
        
    elif filename_lower.endswith('.docx') or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        logger.info(f"Parsing DOCX resume: {filename}")
        return extract_text_from_docx(file_content)
        
    elif filename_lower.endswith('.doc'):
        # Legacy .doc format - not directly supported
        logger.warning(f"Legacy .doc format detected: {filename}")
        raise ValueError("Legacy .doc format not supported. Please convert to .docx or .pdf")
        
    elif filename_lower.endswith(('.txt', '.md', '.markdown')):
        logger.info(f"Parsing text resume: {filename}")
        return extract_text_from_txt(file_content)
        
    else:
        # Try to parse as text
        logger.info(f"Unknown format, attempting text extraction: {filename}")
        try:
            return extract_text_from_txt(file_content)
        except Exception:
            raise ValueError(f"Unsupported file format: {filename}")
